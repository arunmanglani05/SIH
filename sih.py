from flask import Flask, render_template, request, redirect, url_for, session,jsonify
from flask_mysqldb import MySQL
import MySQLdb.cursors
import re
import pickle
import pandas as pd
import xgboost as xgb
from sklearn.model_selection import train_test_split
import warnings
import numpy as np
import random
from random import choice, sample
import PyPDF2
from docx import Document

model = pickle.load(open("C:\\Users\\Dell\\complete web development\\sih_2024\\sih1.pkl", "rb"))
  
  
app = Flask(__name__)
  
  
app.secret_key = 'xyzsdfg'
  
app.config['MYSQL_HOST'] = 'localhost'
app.config['MYSQL_USER'] = 'root'
app.config['MYSQL_PASSWORD'] = ''
app.config['MYSQL_DB'] = 'sih'
  
mysql = MySQL(app)
  
@app.route('/')
@app.route('/login')
def login():
    return render_template('home.html')

# @app.route('/test', methods =['GET', 'POST'])
# def test():
#     data1=request.files['pdf']

#     # pip install PyPDF2 python-docx nltk spacy scikit-learn
#     def pdf_to_text(pdf_file):
#         with open(pdf_file, 'rb') as file:
#             reader = PyPDF2.PdfReader(file)
#             text = ''
#             for page in reader.pages:
#                 text += page.extract_text()
#         return text
#     from docx import Document
#     def docx_to_text(docx_file):
#         doc = Document(docx_file)
#         return '\n'.join([para.text for para in doc.paragraphs])
#     def extract_skills_section(text):
#         lines = text.splitlines()
#         skills_section = ""
#         capture = False

#         for line in lines:
#             if "skills" in line.lower():
#                 capture = True
#             elif capture:
#                 if line.strip() == "" or any(header in line.lower() for header in ["experience", "education", "work"]):
#                     break
#                 skills_section += line + " "

#         return skills_section.strip()
#     def extract_skills(skills_section):
#         skills = skills_section.split(",")  # Assume skills are comma-separated
#         return [skill.strip() for skill in skills if skill.strip()]
    
#     def process_resume(file):
#         if file.endswith('.pdf'):
#             text = pdf_to_text(file)
#         elif file.endswith('.docx'):
#             text = docx_to_text(file)
#         else:
#             raise ValueError("Unsupported file type")

#         skills_section = extract_skills_section(text)
#         skills = extract_skills(skills_section)
#         skills_string = ' '.join(skills)
# # pr    int(skills_string)

#         return skills_string
    
#     skills=process_resume(data1)
#     return render_template('after.html', s=skills)

@app.route('/test', methods=['POST'])
def test():
    if 'pdf' not in request.files:
        # flash('No file part in the request')
        return redirect(url_for('upload_page'))  # Redirect to an upload page

    file = request.files['pdf']
    pos=request.form['position']
    try:
        conn = mysql.connect  # Assuming you have mysql configured
        if conn:
            cursor = conn.cursor(MySQLdb.cursors.DictCursor)
            query = "SELECT skills FROM openings WHERE position = %s"
            cursor.execute(query, (pos,))
            results = cursor.fetchone()
            # results = cursor.fetchall()
            cursor.close()
            conn.commit()
            if results and 'skills' in results:
                sk = results['skills']
            else:
                sk = ""  # Handle the case where no matching position is found
        else:
            return "Connection not established"
    except Exception as e:
            return f"An error occurred: {str(e)}"
    
    if file.filename == '':
        # flash('No selected file')
        return redirect(url_for('upload_page'))

    if file:
        # Get the file extension
        file_extension = file.filename.rsplit('.', 1)[1].lower()

        def pdf_to_text(pdf_file):
            reader = PyPDF2.PdfReader(pdf_file)
            text = ''
            for page in reader.pages:
                text += page.extract_text()
            return text

        def docx_to_text(docx_file):
            doc = Document(docx_file)
            return '\n'.join([para.text for para in doc.paragraphs])

        def extract_skills_section(text):
            lines = text.splitlines()
            skills_section = ""
            capture = False

            for line in lines:
                if "skills" in line.lower():
                    capture = True
                elif capture:
                    if line.strip() == "" or any(header in line.lower() for header in ["experience", "education", "work"]):
                        break
                    skills_section += line + " "

            return skills_section.strip()

        def extract_skills(skills_section):
            skills = skills_section.split(",")  # Assume skills are comma-separated
            return [skill.strip() for skill in skills if skill.strip()]

        def process_resume(file):
            if file_extension == 'pdf':
                text = pdf_to_text(file)
            elif file_extension == 'docx':
                text = docx_to_text(file)
            else:
                raise ValueError("Unsupported file type")

            skills_section = extract_skills_section(text)
            skills = extract_skills(skills_section)
            skills_string = ' '.join(skills)
            return skills_string
        
        skills = process_resume(file)
        def replace_non_letters_with_comma(input_string):
            # Replace all characters that are not letters (a-z or A-Z) with a comma
            return re.sub(r'[^a-zA-Z+]', ',', input_string)
        skills=replace_non_letters_with_comma(skills)
        def find_common_skills(text1, text2):
            # Split the texts by commas and convert to sets
            skills1 = set(text1.split(','))
            skills2 = set(text2.split(','))

            # Find common skills
            common_skills = skills1.intersection(skills2)

            # Join the common skills into a string separated by commas
            return ','.join(common_skills)
        skills1=find_common_skills(skills,sk)
        return render_template('home.html', s=skills1,s1=skills,s2=sk)
    
@app.route('/match_skills', methods=['POST'])
# def match_skills():
#     def get_matching_professions(input_skills):
#         input_skills_set = set(input_skills.lower().split(","))
        
#         try:
#             conn = mysql.connect  # Assuming you have mysql configured
#             if conn:
#                 cursor = conn.cursor(MySQLdb.cursors.DictCursor)
#                 cursor.execute("SELECT professor, skills FROM prof")
#                 results = cursor.fetchall()
#                 cursor.close()
#                 conn.commit()
#             else:
#                 return "Connection not established"

#         except Exception as e:
#             return f"An error occurred: {str(e)}"

#         matches = []
#         for row in results:
#             profession = row['professor']
#             skills = row['skills']
#             profession_skills_set = set(skills.lower().split(","))
#             match_score = len(input_skills_set.intersection(profession_skills_set))
#             if match_score > 0:
#                 matches.append((profession, match_score))

#         matches.sort(key=lambda x: x[1], reverse=True)  # Sort by match score
        
#         # Convert the top 5 matches into a concatenated string
#         top_matches = " ".join([f"{profession}({score})" for profession, score in matches[:5]])
        
#         return top_matches

#     # def get_matching_professions(input_skills):
#     #     input_skills_set = set(input_skills.lower().split(","))
#     #     # conn = sqlite3.connect('database.db')  # Connect to your database
#     #     # cursor = conn.cursor()

#     #     # cursor.execute("SELECT profession_name, skills FROM prof")
#     #     # results = cursor.fetchall()
#     #     try:
#     #         conn = mysql.connect
#     #         if conn:
#     #                 cursor = conn.cursor(MySQLdb.cursors.DictCursor)
#     #                 # cursor = mysql.connect.cursor(MySQLdb.cursors.DictCursor)
#     #                 cursor.execute("SELECT profession_name, skills FROM prof")
#     #                 results = cursor.fetchall()
#     #                 conn.commit()
#     #                 cursor.close()
#     #                 # return render_template("success.html")
#     #                 # message="success"
#     #         else:
#     #             return "connection not extablished"
    
#     #     except Exception as e:
#     #         return f"An error occurred: {str(e)}"

#     #     matches = []
#     #     for profession, skills in results:
#     #         profession_skills_set = set(skills.lower().split(","))
#     #         match_score = len(input_skills_set.intersection(profession_skills_set))
#     #         if match_score > 0:
#     #             matches.append((profession, match_score))

#     #     matches.sort(key=lambda x: x[1], reverse=True)  # Sort by match score
#     #     return matches[:5]  # Return top 5

#     input_skills = request.form['skills']
#     top_matches = get_matching_professions(input_skills)
#     return render_template('after1.html', s=top_matches)
# def match_skills():
#     def get_matching_professions(input_skills):
#         input_skills_set = set(input_skills.lower().split(","))

#         try:
#             conn = mysql.connect  # Assuming you have mysql configured
#             if conn:
#                 cursor = conn.cursor(MySQLdb.cursors.DictCursor)
#                 cursor.execute("SELECT professor, skills FROM prof")
#                 results = cursor.fetchall()
#                 cursor.close()
#                 conn.commit()
#             else:
#                 return "Connection not established"
        
#         except Exception as e:
#             return f"An error occurred: {str(e)}"

#         matches = []
#         for row in results:
#             profession = row['professor']
#             skills = row['skills']
#             profession_skills_set = set(skills.lower().split(","))
#             matched_skills = input_skills_set.intersection(profession_skills_set)
#             match_score = len(matched_skills)
#             if match_score > 0:
#                 matches.append((profession, match_score, matched_skills))

#         # Sort matches by score in descending order
#         matches.sort(key=lambda x: x[1], reverse=True)
        
#         # Format the top 5 matches
#         top_matches = " ".join([f"{profession}({score}) - Matched skills: {', '.join(matched_skills)}" 
#                                 for profession, score, matched_skills in matches[:5]])
        
#         return top_matches

#     input_skills = request.form['skills']
#     top_matches = get_matching_professions(input_skills)
#     return render_template('after1.html', s=top_matches)
def match_skills():
    def get_matching_professions(input_skills):
        input_skills_set = set(input_skills.lower().split(","))

        try:
            conn = mysql.connect  # Assuming you have mysql configured
            if conn:
                cursor = conn.cursor(MySQLdb.cursors.DictCursor)
                cursor.execute("SELECT professor, skills FROM prof")
                results = cursor.fetchall()
                cursor.close()
                conn.commit()
            else:
                return "Connection not established"
        
        except Exception as e:
            return f"An error occurred: {str(e)}"

        matches = []
        for row in results:
            profession = row['professor']
            skills = row['skills']
            profession_skills_set = set(skills.lower().split(","))
            matched_skills = input_skills_set.intersection(profession_skills_set)
            match_score = len(matched_skills)
            if match_score > 0:
                matches.append((profession, match_score, matched_skills))

        # Sort matches by score in descending order
        matches.sort(key=lambda x: x[1], reverse=True)
        
        # Format the top 5 matches
        top_matches = "<br>".join([f"{profession} - Score: {score} - Matched skills: {', '.join(matched_skills)}" 
                                   for profession, score, matched_skills in matches[:5]])
        
        return top_matches

    input_skills = request.form['skills']
    top_matches = get_matching_professions(input_skills)
    return render_template('after1.html', s=top_matches)








if __name__ == "__main__":
    app.run(debug=True)